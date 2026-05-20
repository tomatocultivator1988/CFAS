#!/usr/bin/env python3
"""
Fix Cap_set...F.docx formatting - put each choice on its own line
"""
from docx import Document
import re

# Read the original file
doc = Document('Exam-Main/Cap_set...F.docx')

# Create new document
new_doc = Document()
new_doc.add_heading('CapSet F - Fisheries Examination', 0)

# Extract all paragraphs
for para in doc.paragraphs:
    text = para.text.strip()
    
    if not text:
        continue
    
    # Check if this is a choice line (contains multiple choices)
    # Pattern: "a. text    b. text" or "a. text\tb. text"
    if re.search(r'[a-d]\.\s+.+\s{2,}[a-d]\.\s+', text) or '\t' in text:
        # Split by tabs or multiple spaces before choice letters
        parts = re.split(r'[\t]+|(?=\s{2,}[a-d]\.\s+)', text)
        
        for part in parts:
            part = part.strip()
            if part:
                new_doc.add_paragraph(part)
    else:
        # Regular line - add as is
        new_doc.add_paragraph(text)

# Save the reformatted document
new_doc.save('Exam-Main/Cap_set_F_Formatted.docx')

print("✅ Reformatted file created: Exam-Main/Cap_set_F_Formatted.docx")
print("")
print("📤 Upload this file to the admin panel!")
print("🎯 The AI will read the answer key table at the end automatically")
