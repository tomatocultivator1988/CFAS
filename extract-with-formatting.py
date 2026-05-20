#!/usr/bin/env python3
"""
DOCX Text Extractor with Formatting Detection
Extracts text and marks BOLD and HIGHLIGHTED text with special markers
"""
import sys
import os
from docx import Document
from docx.shared import RGBColor

def has_highlight(run):
    """Check if run has highlight/background color"""
    try:
        if run.font.highlight_color:
            return True
    except:
        pass
    return False

def has_background_color(run):
    """Check if run has background shading"""
    try:
        if run._element.rPr is not None:
            shd = run._element.rPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
            if shd is not None:
                fill = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                if fill and fill != 'auto' and fill != 'FFFFFF':
                    return True
    except:
        pass
    return False

def extract_text_with_formatting(docx_path, output_path):
    """Extract text and mark bold/highlighted text"""
    try:
        doc = Document(docx_path)
        all_text = []
        
        # Extract from paragraphs with formatting
        for para in doc.paragraphs:
            para_text = []
            
            for run in para.runs:
                text = run.text
                if not text:
                    continue
                
                # Check if bold
                is_bold = run.bold == True
                
                # Check if highlighted or has background color
                is_highlighted = has_highlight(run) or has_background_color(run)
                
                # Mark formatted text with special markers
                if is_bold or is_highlighted:
                    # Use markers that AI can easily recognize
                    text = f"**{text}**"  # Mark as correct answer
                
                para_text.append(text)
            
            if para_text:
                full_para = ''.join(para_text).strip()
                if full_para:
                    all_text.append(full_para)
        
        # Extract from tables with formatting
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_parts = []
                    for para in cell.paragraphs:
                        para_parts = []
                        for run in para.runs:
                            text = run.text
                            if not text:
                                continue
                            
                            is_bold = run.bold == True
                            is_highlighted = has_highlight(run) or has_background_color(run)
                            
                            if is_bold or is_highlighted:
                                text = f"**{text}**"
                            
                            para_parts.append(text)
                        
                        if para_parts:
                            cell_parts.append(''.join(para_parts))
                    
                    if cell_parts:
                        row_text.append(' '.join(cell_parts).strip())
                
                if row_text:
                    # Use newline instead of tab for better AI parsing
                    # This ensures each cell content is on its own line
                    all_text.append('\n'.join(row_text))
        
        # Write to output file
        full_text = '\n'.join(all_text)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(full_text)
        
        # Count marked items
        marked_count = full_text.count('**')
        print(f"SUCCESS: Extracted {len(full_text)} characters, {len(all_text)} lines, {marked_count//2} marked items")
        return 0
        
    except Exception as e:
        print(f"ERROR: {str(e)}", file=sys.stderr)
        import traceback
        traceback.print_exc(file=sys.stderr)
        return 1

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python extract-with-formatting.py <docx_file> <output_file>", file=sys.stderr)
        sys.exit(1)
    
    docx_file = sys.argv[1]
    output_file = sys.argv[2]
    
    if not os.path.exists(docx_file):
        print(f"ERROR: File not found: {docx_file}", file=sys.stderr)
        sys.exit(1)
    
    exit_code = extract_text_with_formatting(docx_file, output_file)
    sys.exit(exit_code)
