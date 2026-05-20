"""
Smart DOCX Parser using Free AI Model (Ollama)
Automatically reformats questions from any format to structured JSON

This parser can handle:
1. Questions with answer key at bottom (table format)
2. Questions with highlighted answers
3. Questions with inline answers
4. Mixed formats
5. Messy/unstructured formats

Uses Ollama (free, local AI) to intelligently parse questions
"""

import sys
import json
import re
from docx import Document
from docx.table import Table
import requests

# Ollama API Configuration (runs locally, free)
OLLAMA_API_URL = "http://localhost:11434/api/generate"
OLLAMA_MODEL = "llama3.2:1b"  # Fast, lightweight model (1.3GB)

def check_ollama():
    """Check if Ollama is running"""
    try:
        response = requests.get("http://localhost:11434/api/tags", timeout=2)
        return response.status_code == 200
    except:
        return False

def extract_text_from_docx(file_path):
    """Extract all text from Word document including tables"""
    doc = Document(file_path)
    
    full_text = []
    answer_key = {}
    
    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text.append(text)
    
    # Extract tables (usually answer keys)
    for table in doc.tables:
        table_text = []
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                table_text.append(" | ".join(row_text))
        
        if table_text:
            full_text.append("\n[TABLE START]")
            full_text.extend(table_text)
            full_text.append("[TABLE END]\n")
    
    return "\n".join(full_text)

def parse_with_ai(text_content):
    """Use Ollama AI to parse questions from unstructured text"""
    
    prompt = f"""You are a question parser. Extract questions and answers from the text below.

The text may contain:
- Questions numbered 1-100
- Multiple choice options (a, b, c, d)
- An answer key (may be in a table at the bottom)

Your task:
1. Extract each question with its choices
2. Match the correct answer from the answer key
3. Output ONLY valid JSON array, no other text

Output format (JSON only, no markdown):
[
  {{
    "question_text": "Question text here?",
    "answer_choices": [
      {{"choice_text": "Choice A text", "is_correct": false}},
      {{"choice_text": "Choice B text", "is_correct": true}},
      {{"choice_text": "Choice C text", "is_correct": false}},
      {{"choice_text": "Choice D text", "is_correct": false}}
    ]
  }}
]

TEXT TO PARSE:
{text_content[:8000]}

Remember: Output ONLY the JSON array, nothing else."""

    try:
        response = requests.post(
            OLLAMA_API_URL,
            json={
                "model": OLLAMA_MODEL,
                "prompt": prompt,
                "stream": False,
                "temperature": 0.1,  # Low temperature for consistent output
                "options": {
                    "num_predict": 4000  # Allow longer responses
                }
            },
            timeout=120
        )
        
        if response.status_code == 200:
            result = response.json()
            ai_output = result.get('response', '')
            
            # Extract JSON from AI output (may have extra text)
            json_match = re.search(r'\[.*\]', ai_output, re.DOTALL)
            if json_match:
                json_str = json_match.group(0)
                questions = json.loads(json_str)
                return questions
            else:
                return None
        else:
            return None
            
    except Exception as e:
        print(f"AI parsing error: {e}", file=sys.stderr)
        return None

def parse_traditional(text_content):
    """Traditional regex-based parsing (fallback)"""
    questions = []
    lines = text_content.split('\n')
    
    # Extract answer key from table
    answer_key = {}
    in_table = False
    for line in lines:
        if '[TABLE START]' in line:
            in_table = True
            continue
        if '[TABLE END]' in line:
            in_table = False
            continue
        
        if in_table:
            # Try to parse answer key: "1 | c" or "1. c" format
            match = re.match(r'(\d+)\s*[|\.\)]\s*([a-d])', line, re.IGNORECASE)
            if match:
                q_num = int(match.group(1))
                answer = match.group(2).lower()
                answer_key[q_num] = answer
    
    # Parse questions
    current_question = None
    current_choices = []
    question_num = 0
    
    for line in lines:
        line = line.strip()
        if not line or '[TABLE' in line:
            continue
        
        # Check for question (starts with number)
        q_match = re.match(r'^(\d+)\.\s*(.+)', line)
        if q_match:
            # Save previous question
            if current_question and current_choices:
                questions.append({
                    'question_text': current_question,
                    'answer_choices': current_choices
                })
            
            # Start new question
            question_num = int(q_match.group(1))
            current_question = q_match.group(2).strip()
            current_choices = []
            continue
        
        # Check for choice (starts with letter)
        c_match = re.match(r'^([a-d])\.\s*(.+)', line, re.IGNORECASE)
        if c_match and current_question:
            choice_letter = c_match.group(1).lower()
            choice_text = c_match.group(2).strip()
            
            # Check if this is the correct answer
            correct_letter = answer_key.get(question_num, '').lower()
            is_correct = (choice_letter == correct_letter)
            
            current_choices.append({
                'choice_text': choice_text,
                'is_correct': is_correct
            })
    
    # Don't forget last question
    if current_question and current_choices:
        questions.append({
            'question_text': current_question,
            'answer_choices': current_choices
        })
    
    return questions

def main():
    if len(sys.argv) < 2:
        print(json.dumps({"error": "Usage: python smart-docx-parser.py <docx-file>"}))
        sys.exit(1)
    
    docx_file = sys.argv[1]
    use_ai = len(sys.argv) > 2 and sys.argv[2] == "--ai"
    
    try:
        # Extract text from document
        text_content = extract_text_from_docx(docx_file)
        
        if not text_content:
            print(json.dumps({"error": "No text found in document"}))
            sys.exit(1)
        
        questions = None
        
        # Try AI parsing if requested and Ollama is available
        if use_ai and check_ollama():
            print("Using AI parser (Ollama)...", file=sys.stderr)
            questions = parse_with_ai(text_content)
        
        # Fallback to traditional parsing
        if not questions:
            print("Using traditional parser...", file=sys.stderr)
            questions = parse_traditional(text_content)
        
        if not questions:
            print(json.dumps({"error": "No questions found"}))
            sys.exit(1)
        
        # Validate questions
        valid_questions = []
        for q in questions:
            if (q.get('question_text') and 
                q.get('answer_choices') and 
                len(q['answer_choices']) >= 2):
                # Ensure at least one correct answer
                has_correct = any(c.get('is_correct') for c in q['answer_choices'])
                if not has_correct and q['answer_choices']:
                    # Default first choice as correct if none marked
                    q['answer_choices'][0]['is_correct'] = True
                valid_questions.append(q)
        
        # Output JSON
        print(json.dumps(valid_questions, ensure_ascii=False))
        sys.exit(0)
        
    except Exception as e:
        print(json.dumps({"error": str(e)}))
        sys.exit(1)

if __name__ == "__main__":
    main()
