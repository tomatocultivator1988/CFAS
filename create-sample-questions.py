"""
Create a sample Word document with questions
Correct answers are highlighted in yellow
"""

from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

# Create document
doc = Document()

# Add title
doc.add_heading('Sample Questions', 0)

# Questions data
questions = [
    {
        'question': 'What is the capital of India?',
        'choices': ['Mumbai', 'Delhi', 'Chennai', 'Kolkata'],
        'correct': 1  # Delhi (index 1)
    },
    {
        'question': 'Which animal is known as the King of the Jungle?',
        'choices': ['Tiger', 'Elephant', 'Lion', 'Leopard'],
        'correct': 2  # Lion
    },
    {
        'question': 'How many days are there in a leap year?',
        'choices': ['364', '365', '366', '367'],
        'correct': 2  # 366
    },
    {
        'question': 'Which planet is closest to the Sun?',
        'choices': ['Venus', 'Earth', 'Mercury', 'Mars'],
        'correct': 2  # Mercury
    },
    {
        'question': 'Who invented the telephone?',
        'choices': ['Thomas Edison', 'Alexander Graham Bell', 'Isaac Newton', 'Albert Einstein'],
        'correct': 1  # Alexander Graham Bell
    },
    {
        'question': 'What is the national flower of India?',
        'choices': ['Rose', 'Lotus', 'Sunflower', 'Lily'],
        'correct': 1  # Lotus
    },
    {
        'question': 'Which is the largest continent in the world?',
        'choices': ['Africa', 'Europe', 'Asia', 'Australia'],
        'correct': 2  # Asia
    },
    {
        'question': 'What do we call a house of books?',
        'choices': ['School', 'Museum', 'Library', 'Office'],
        'correct': 2  # Library
    },
    {
        'question': 'Which gas is essential for breathing?',
        'choices': ['Carbon dioxide', 'Hydrogen', 'Nitrogen', 'Oxygen'],
        'correct': 3  # Oxygen
    },
    {
        'question': 'How many letters are there in the English alphabet?',
        'choices': ['24', '25', '26', '27'],
        'correct': 2  # 26
    }
]

# Add questions
for i, q in enumerate(questions, 1):
    # Add question
    para = doc.add_paragraph()
    para.add_run(f"{i}. {q['question']}")
    
    # Add choices
    letters = ['a', 'b', 'c', 'd']
    for j, choice in enumerate(q['choices']):
        para = doc.add_paragraph()
        run = para.add_run(f"{letters[j]}. {choice}")
        
        # Highlight correct answer
        if j == q['correct']:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    
    # Add blank line
    doc.add_paragraph()

# Save document
doc.save('Exam-Main/Sample_Questions_Highlighted.docx')
print("✅ Created Sample_Questions_Highlighted.docx with highlighted correct answers")
print("📝 Correct answers are highlighted in YELLOW")
print("📤 You can now import this file using the Question Management page")
