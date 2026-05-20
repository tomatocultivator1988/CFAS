from docx import Document
import re

doc = Document('Aquaculture_Reviewer-1_Shellcheck-Full.docx')

print("=" * 60)
print("SHELLCHECK FILE ANALYSIS")
print("=" * 60)

# Count paragraphs
total_paras = len([p for p in doc.paragraphs if p.text.strip()])
print(f"\nTotal paragraphs: {total_paras}")

# Find bold text
bold_count = 0
bold_samples = []
for i, p in enumerate(doc.paragraphs[:200]):
    for run in p.runs:
        if run.bold and run.text.strip():
            bold_count += 1
            if len(bold_samples) < 10:
                bold_samples.append(f"Para {i+1}: {run.text[:60]}")

print(f"\nBold text found (first 200 paras): {bold_count}")
print("\n=== BOLD TEXT SAMPLES ===")
for sample in bold_samples:
    print(sample)

# Check for question pattern
text = '\n'.join([p.text for p in doc.paragraphs])
question_pattern = re.findall(r'^\d+\.\s+.{10,}', text, re.MULTILINE)
print(f"\n\nQuestion-like patterns found: {len(question_pattern)}")
print("\n=== FIRST 10 QUESTIONS ===")
for i, q in enumerate(question_pattern[:10]):
    print(f"{i+1}. {q[:80]}")

# Check for tables
print(f"\n\nTables found: {len(doc.tables)}")
if doc.tables:
    print("\n=== FIRST TABLE ===")
    table = doc.tables[0]
    print(f"Rows: {len(table.rows)}, Columns: {len(table.columns)}")
    for i, row in enumerate(table.rows[:3]):
        print(f"Row {i}: {[cell.text[:30] for cell in row.cells]}")

# Check structure
print("\n\n=== DOCUMENT STRUCTURE (First 30 lines) ===")
for i, p in enumerate(doc.paragraphs[:30]):
    if p.text.strip():
        has_bold = any(run.bold for run in p.runs)
        bold_marker = " [BOLD]" if has_bold else ""
        print(f"{i+1}. {p.text[:80]}{bold_marker}")
