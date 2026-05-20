#!/usr/bin/env python3
"""
Format Cap_set...E.docx with proper formatting and accurate answer key table.
- Splits choices onto separate lines
- Adds complete answer key table at the end
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# EXACT answer key from Cap_set...E.docx table
ANSWER_KEY = {
    1: 'c', 2: 'd', 3: 'b', 4: 'a', 5: 'b', 6: 'a', 7: 'd', 8: 'c', 9: 'a', 10: 'c',
    11: 'b', 12: 'd', 13: 'b', 14: 'c', 15: 'a', 16: 'b', 17: 'c', 18: 'b', 19: 'c', 20: 'd',
    21: 'b', 22: 'a', 23: 'd', 24: 'c', 25: 'c', 26: 'd', 27: 'b', 28: 'b', 29: 'c', 30: 'a',
    31: 'd', 32: 'b', 33: 'd', 34: 'a', 35: 'd', 36: 'c', 37: 'b', 38: 'c', 39: 'd', 40: 'd',
    41: 'a', 42: 'c', 43: 'a', 44: 'd', 45: 'd', 46: 'c', 47: 'a', 48: 'd', 49: 'c', 50: 'd',
    51: 'c', 52: 'd', 53: 'd', 54: 'a', 55: 'c', 56: 'b', 57: 'd', 58: 'c', 59: 'd', 60: 'a',
    61: 'b', 62: 'c', 63: 'd', 64: 'c', 65: 'd', 66: 'c', 67: 'b', 68: 'd', 69: 'd', 70: 'b',
    71: 'd', 72: 'a', 73: 'b', 74: 'c', 75: 'a', 76: 'b', 77: 'c', 78: 'd', 79: 'a', 80: 'd',
    81: 'd', 82: 'b', 83: 'd', 84: 'a', 85: 'b', 86: 'c', 87: 'b', 88: 'a', 89: 'b', 90: 'a',
    91: 'b', 92: 'c', 93: 'd', 94: 'a', 95: 'b', 96: 'a', 97: 'b', 98: 'c', 99: 'd', 100: 'b'
}

def split_choices_on_tabs(text):
    """Split choices that are on the same line with tabs/spaces between them."""
    # Replace multiple tabs/spaces with newline
    text = re.sub(r'\t+', '\n', text)
    text = re.sub(r'  +', '\n', text)
    return text

def process_document(input_path, output_path):
    """Process the DOCX file and create formatted version with answer key."""
    print(f"📖 Reading: {input_path}")
    doc = Document(input_path)
    new_doc = Document()
    
    # Add title
    title = new_doc.add_paragraph("CapSet_E")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)
    new_doc.add_paragraph()
    
    # Process each paragraph
    question_count = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Skip empty paragraphs and answer key section
        if not text or 'Answer Key' in text or 'CapSet' in text:
            continue
        
        # Check if it's a question (starts with number)
        if re.match(r'^\d+\.', text):
            question_count += 1
            # Add question
            p = new_doc.add_paragraph(text)
            p.runs[0].font.size = Pt(11)
        else:
            # It's a choice line - split if needed
            formatted_text = split_choices_on_tabs(text)
            lines = [line.strip() for line in formatted_text.split('\n') if line.strip()]
            
            for line in lines:
                # Add each choice on its own line
                p = new_doc.add_paragraph(line)
                p.runs[0].font.size = Pt(11)
    
    print(f"✅ Processed {question_count} questions")
    
    # Add answer key section
    new_doc.add_page_break()
    title = new_doc.add_paragraph("Answer Key")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)
    new_doc.add_paragraph()
    
    # Create answer key table (4 columns x 25 rows)
    table = new_doc.add_table(rows=26, cols=4)
    table.style = 'Table Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    for i, col_num in enumerate([1, 26, 51, 76]):
        header_cells[i].text = f"Q{col_num}-{col_num+24}"
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Fill answer key
    for row_idx in range(1, 26):
        for col_idx in range(4):
            q_num = row_idx + (col_idx * 25)
            if q_num <= 100:
                answer = ANSWER_KEY.get(q_num, '?')
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = f"{q_num}. {answer}"
                cell.paragraphs[0].runs[0].font.size = Pt(10)
    
    # Save
    print(f"💾 Saving: {output_path}")
    new_doc.save(output_path)
    print(f"✅ Done! File saved with {question_count} questions and complete answer key.")

if __name__ == "__main__":
    input_file = "Exam-Main/Cap_set...E.docx"
    output_file = "Exam-Main/Cap_set_E_Complete.docx"
    
    try:
        process_document(input_file, output_file)
        print("\n✨ SUCCESS! Cap_set_E_Complete.docx is ready for import.")
        print("📋 Answer key table included with all 100 answers.")
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
