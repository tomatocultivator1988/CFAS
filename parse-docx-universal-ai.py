"""
UNIVERSAL AI-POWERED DOCX PARSER
Uses Gemini AI to understand ANY format and restructure to standard JSON

This parser can handle:
- ANY question format
- ANY choice format  
- ANY answer key format
- Mixed formats
- Messy/unstructured documents

The AI understands the content and restructures it automatically!
"""

import sys
import json
import re
from docx import Document
import requests

# Gemini API Configuration
GEMINI_API_KEY = "[REDACTED_GEMINI_API_KEY]"
GEMINI_API_URL = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-pro:generateContent?key={GEMINI_API_KEY}"

def extract_all_text_from_docx(file_path):
    """Extract ALL text from Word document including tables"""
    doc = Document(file_path)
    
    all_text = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            all_text.append(text)
    
    # Extract tables
    for table in doc.tables:
        all_text.append("\n=== TABLE START ===")
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_data.append(cell_text)
            if row_data:
                all_text.append(" | ".join(row_data))
        all_text.append("=== TABLE END ===\n")
    
    return "\n".join(all_text)

def parse_with_universal_ai(text_content):
    """Use Gemini AI to parse and restructure ANY format to standard JSON"""
    
    # Limit text to avoid token limits
    text_sample = text_content[:15000]  # ~100-150 questions
    
    prompt = f"""You are an expert question parser. Your task is to extract exam questions from the text below and convert them to a standardized JSON format.

The text may contain questions in ANY format:
- Questions may be numbered (1., Q1, Question 1, etc.)
- Choices may be lettered (a., A., a), A), etc.)
- Choices may be inline (on same line) or separate lines
- Answer key may be in a table, list, or inline
- Format may be messy or inconsistent

YOUR TASK:
1. Identify ALL questions and their choices
2. Determine the CORRECT answer for each question (from answer key or context)
3. Output ONLY a valid JSON array in this EXACT format:

[
  {{
    "question_text": "Full question text here",
    "answer_choices": [
      {{"choice_text": "First choice text", "is_correct": true}},
      {{"choice_text": "Second choice text", "is_correct": false}},
      {{"choice_text": "Third choice text", "is_correct": false}},
      {{"choice_text": "Fourth choice text", "is_correct": false}}
    ]
  }},
  {{
    "question_text": "Next question text",
    "answer_choices": [
      {{"choice_text": "Choice A", "is_correct": false}},
      {{"choice_text": "Choice B", "is_correct": true}},
      {{"choice_text": "Choice C", "is_correct": false}},
      {{"choice_text": "Choice D", "is_correct": false}}
    ]
  }}
]

CRITICAL RULES:
1. Output ONLY the JSON array - no markdown, no explanations, no extra text
2. Each question must have "question_text" and "answer_choices"
3. Each choice must have "choice_text" and "is_correct" (boolean)
4. EXACTLY ONE choice per question must have "is_correct": true
5. Extract the FULL question text (not just the number)
6. Extract the FULL choice text (not just the letter)
7. If answer key is unclear, use your knowledge to determine correct answer
8. Clean up any formatting issues (tabs, extra spaces, etc.)

TEXT TO PARSE:
{text_sample}

Remember: Output ONLY the JSON array. Start with [ and end with ]"""

    try:
        print("Sending to Gemini AI...", file=sys.stderr)
        response = requests.post(
            GEMINI_API_URL,
            json={
                "contents": [{
                    "parts": [{
                        "text": prompt
                    }]
                }],
                "generationConfig": {
                    "temperature": 0.1,  # Low temperature for consistency
                    "maxOutputTokens": 8000,
                    "topP": 0.8,
                    "topK": 10
                }
            },
            timeout=90  # Longer timeout for complex parsing
        )
        
        if response.status_code == 200:
            result = response.json()
            
            if 'candidates' in result and len(result['candidates']) > 0:
                candidate = result['candidates'][0]
                if 'content' in candidate and 'parts' in candidate['content']:
                    ai_output = candidate['content']['parts'][0]['text']
                    
                    print("AI response received, parsing JSON...", file=sys.stderr)
                    
                    # Clean up the output
                    ai_output = ai_output.strip()
                    
                    # Remove markdown code blocks if present
                    ai_output = re.sub(r'```json\s*', '', ai_output)
                    ai_output = re.sub(r'```\s*', '', ai_output)
                    
                    # Try to parse JSON directly
                    try:
                        questions = json.loads(ai_output)
                        return questions
                    except json.JSONDecodeError as e:
                        print(f"JSON parse error: {e}", file=sys.stderr)
                        
                        # Try to extract JSON array from text
                        json_match = re.search(r'\[.*\]', ai_output, re.DOTALL)
                        if json_match:
                            json_str = json_match.group(0)
                            questions = json.loads(json_str)
                            return questions
                        else:
                            print(f"Could not find JSON in output", file=sys.stderr)
                            print(f"AI output (first 500 chars): {ai_output[:500]}", file=sys.stderr)
                            return None
            
            print(f"Unexpected API response structure", file=sys.stderr)
            return None
        else:
            print(f"API error: {response.status_code}", file=sys.stderr)
            print(f"Response: {response.text[:500]}", file=sys.stderr)
            return None
            
    except Exception as e:
        print(f"Error during AI parsing: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc(file=sys.stderr)
        return None

def validate_and_fix_questions(questions):
    """Validate and fix common issues in parsed questions"""
    if not questions or not isinstance(questions, list):
        return []
    
    valid_questions = []
    
    for i, q in enumerate(questions, 1):
        # Must have question text
        if not q.get('question_text'):
            print(f"Warning: Question {i} has no text, skipping", file=sys.stderr)
            continue
        
        # Must have answer choices
        if not q.get('answer_choices') or not isinstance(q['answer_choices'], list):
            print(f"Warning: Question {i} has no choices, skipping", file=sys.stderr)
            continue
        
        # Must have at least 2 choices
        if len(q['answer_choices']) < 2:
            print(f"Warning: Question {i} has less than 2 choices, skipping", file=sys.stderr)
            continue
        
        # Ensure exactly one correct answer
        correct_count = sum(1 for c in q['answer_choices'] if c.get('is_correct'))
        
        if correct_count == 0:
            # No correct answer - default first choice
            print(f"Warning: Question {i} has no correct answer, defaulting to first choice", file=sys.stderr)
            q['answer_choices'][0]['is_correct'] = True
        elif correct_count > 1:
            # Multiple correct answers - keep only first
            print(f"Warning: Question {i} has multiple correct answers, keeping only first", file=sys.stderr)
            found_first = False
            for c in q['answer_choices']:
                if c.get('is_correct'):
                    if found_first:
                        c['is_correct'] = False
                    else:
                        found_first = True
        
        # Clean up choice text
        for c in q['answer_choices']:
            if 'choice_text' in c:
                c['choice_text'] = c['choice_text'].strip()
        
        valid_questions.append(q)
    
    return valid_questions

def main():
    if len(sys.argv) < 2:
        print(json.dumps({"error": "Usage: python parse-docx-universal-ai.py <docx-file>"}))
        sys.exit(1)
    
    docx_file = sys.argv[1]
    
    try:
        # Step 1: Extract all text from document
        print(f"Extracting text from: {docx_file}", file=sys.stderr)
        text_content = extract_all_text_from_docx(docx_file)
        
        if not text_content or len(text_content) < 50:
            print(json.dumps({"error": "Document is empty or too short"}))
            sys.exit(1)
        
        print(f"Extracted {len(text_content)} characters", file=sys.stderr)
        
        # Step 2: Use AI to parse and restructure
        print("Using Gemini AI to parse and restructure...", file=sys.stderr)
        questions = parse_with_universal_ai(text_content)
        
        if not questions:
            print(json.dumps({"error": "AI failed to parse questions"}))
            sys.exit(1)
        
        print(f"AI parsed {len(questions)} questions", file=sys.stderr)
        
        # Step 3: Validate and fix
        print("Validating questions...", file=sys.stderr)
        valid_questions = validate_and_fix_questions(questions)
        
        if not valid_questions:
            print(json.dumps({"error": "No valid questions found after validation"}))
            sys.exit(1)
        
        print(f"Validated {len(valid_questions)} questions", file=sys.stderr)
        
        # Step 4: Output JSON
        print(json.dumps(valid_questions, ensure_ascii=False, indent=2))
        sys.exit(0)
        
    except Exception as e:
        print(json.dumps({"error": str(e)}))
        import traceback
        traceback.print_exc(file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()
