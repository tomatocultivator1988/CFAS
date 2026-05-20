"""
Smart DOCX Parser using Google Gemini AI (FREE)
Automatically reformats questions from any format to structured JSON

Uses Gemini 1.5 Flash (FREE tier: 15 requests/minute, 1 million tokens/day)
No installation required - just needs API key!
"""

import sys
import json
import re
from docx import Document
import requests

# Gemini API Configuration
GEMINI_API_KEY = "[REDACTED_GEMINI_API_KEY]"
GEMINI_API_URL = f"https://generativelanguage.googleapis.com/v1/models/gemini-1.5-flash:generateContent?key={GEMINI_API_KEY}"

def extract_text_from_docx(file_path):
    """Extract all text from Word document including tables"""
    doc = Document(file_path)
    
    full_text = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text.append(text)
    
    # Extract tables (usually answer keys)
    for table in doc.tables:
        full_text.append("\n[ANSWER KEY TABLE]")
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                full_text.append(" | ".join(row_text))
        full_text.append("[END TABLE]\n")
    
    return "\n".join(full_text)

def parse_with_gemini(text_content):
    """Use Gemini AI to parse questions from any format"""
    
    # Limit text to avoid token limits (first 10000 chars should cover ~100 questions)
    text_sample = text_content[:10000]
    
    prompt = f"""You are a question parser. Extract questions and multiple choice answers from the text below.

The text contains exam questions that may be in various formats:
- Questions numbered 1-100
- Multiple choice options (a, b, c, d or A, B, C, D)
- An answer key (may be in a table, list, or inline)

Your task:
1. Extract EVERY question with its choices
2. Match the correct answer from the answer key
3. Output ONLY a valid JSON array

Output format (MUST be valid JSON, no markdown, no extra text):
[
  {{
    "question_text": "What is aquaculture?",
    "answer_choices": [
      {{"choice_text": "Farming of fish", "is_correct": true}},
      {{"choice_text": "Farming of crops", "is_correct": false}},
      {{"choice_text": "Farming of livestock", "is_correct": false}},
      {{"choice_text": "Farming of poultry", "is_correct": false}}
    ]
  }}
]

CRITICAL RULES:
- Output ONLY the JSON array
- No markdown code blocks (no ```json)
- No explanations or extra text
- Each question must have "question_text" and "answer_choices"
- Each choice must have "choice_text" and "is_correct" (boolean)
- Exactly ONE choice per question should have "is_correct": true

TEXT TO PARSE:
{text_sample}"""

    try:
        response = requests.post(
            GEMINI_API_URL,
            json={
                "contents": [{
                    "parts": [{
                        "text": prompt
                    }]
                }],
                "generationConfig": {
                    "temperature": 0.1,
                    "maxOutputTokens": 8000
                }
            },
            timeout=60
        )
        
        if response.status_code == 200:
            result = response.json()
            
            # Extract text from Gemini response
            if 'candidates' in result and len(result['candidates']) > 0:
                candidate = result['candidates'][0]
                if 'content' in candidate and 'parts' in candidate['content']:
                    ai_output = candidate['content']['parts'][0]['text']
                    
                    # Clean up the output - remove markdown code blocks if present
                    ai_output = re.sub(r'```json\s*', '', ai_output)
                    ai_output = re.sub(r'```\s*', '', ai_output)
                    ai_output = ai_output.strip()
                    
                    # Try to parse JSON
                    try:
                        questions = json.loads(ai_output)
                        return questions
                    except json.JSONDecodeError:
                        # Try to extract JSON array from text
                        json_match = re.search(r'\[.*\]', ai_output, re.DOTALL)
                        if json_match:
                            json_str = json_match.group(0)
                            questions = json.loads(json_str)
                            return questions
            
            print(f"Gemini API response: {json.dumps(result, indent=2)}", file=sys.stderr)
            return None
        else:
            print(f"Gemini API error: {response.status_code} - {response.text}", file=sys.stderr)
            return None
            
    except Exception as e:
        print(f"Gemini parsing error: {e}", file=sys.stderr)
        return None

def parse_traditional(text_content):
    """Traditional regex-based parsing (fallback)"""
    questions = []
    lines = text_content.split('\n')
    
    # Extract answer key from table
    answer_key = {}
    in_table = False
    for line in lines:
        if '[ANSWER KEY TABLE]' in line:
            in_table = True
            continue
        if '[END TABLE]' in line:
            in_table = False
            continue
        
        if in_table:
            # Try to parse answer key: "1 | c" format
            match = re.match(r'(\d+)\s*\|\s*([a-d])', line, re.IGNORECASE)
            if match:
                q_num = int(match.group(1))
                answer = match.group(2).lower()
                answer_key[q_num] = answer
    
    # Parse questions
    current_question = None
    current_choices = []
    question_num = 0
    
    for line in lines:
        line = line.strip()
        if not line or '[TABLE' in line or '[END' in line:
            continue
        
        # Check for question (starts with number)
        q_match = re.match(r'^(\d+)\.\s*(.+)', line)
        if q_match:
            # Save previous question
            if current_question and current_choices:
                questions.append({
                    'question_text': current_question,
                    'answer_choices': current_choices
                })
            
            # Start new question
            question_num = int(q_match.group(1))
            current_question = q_match.group(2).strip()
            current_choices = []
            continue
        
        # Check for choice (starts with letter)
        c_match = re.match(r'^([a-d])[\.\)]\s*(.+)', line, re.IGNORECASE)
        if c_match and current_question:
            choice_letter = c_match.group(1).lower()
            choice_text = c_match.group(2).strip()
            
            # Check if this is the correct answer
            correct_letter = answer_key.get(question_num, '').lower()
            is_correct = (choice_letter == correct_letter)
            
            current_choices.append({
                'choice_text': choice_text,
                'is_correct': is_correct
            })
    
    # Don't forget last question
    if current_question and current_choices:
        questions.append({
            'question_text': current_question,
            'answer_choices': current_choices
        })
    
    return questions

def validate_questions(questions):
    """Validate and fix questions"""
    valid_questions = []
    
    for q in questions:
        # Must have question text
        if not q.get('question_text'):
            continue
        
        # Must have at least 2 choices
        if not q.get('answer_choices') or len(q['answer_choices']) < 2:
            continue
        
        # Must have at least one correct answer
        has_correct = any(c.get('is_correct') for c in q['answer_choices'])
        if not has_correct:
            # Default first choice as correct
            q['answer_choices'][0]['is_correct'] = True
        
        # Ensure only one correct answer
        correct_count = sum(1 for c in q['answer_choices'] if c.get('is_correct'))
        if correct_count > 1:
            # Keep only first correct answer
            found_first = False
            for c in q['answer_choices']:
                if c.get('is_correct'):
                    if found_first:
                        c['is_correct'] = False
                    else:
                        found_first = True
        
        valid_questions.append(q)
    
    return valid_questions

def main():
    if len(sys.argv) < 2:
        print(json.dumps({"error": "Usage: python parse-docx-gemini.py <docx-file> [--ai]"}))
        sys.exit(1)
    
    docx_file = sys.argv[1]
    use_ai = len(sys.argv) > 2 and sys.argv[2] == "--ai"
    
    try:
        # Extract text from document
        print("Extracting text from document...", file=sys.stderr)
        text_content = extract_text_from_docx(docx_file)
        
        if not text_content:
            print(json.dumps({"error": "No text found in document"}))
            sys.exit(1)
        
        questions = None
        
        # Try AI parsing if requested
        if use_ai:
            print("Using Gemini AI parser...", file=sys.stderr)
            questions = parse_with_gemini(text_content)
            
            if questions:
                print(f"Gemini parsed {len(questions)} questions", file=sys.stderr)
        
        # Fallback to traditional parsing
        if not questions:
            print("Using traditional parser...", file=sys.stderr)
            questions = parse_traditional(text_content)
        
        if not questions:
            print(json.dumps({"error": "No questions found"}))
            sys.exit(1)
        
        # Validate questions
        valid_questions = validate_questions(questions)
        
        print(f"Validated {len(valid_questions)} questions", file=sys.stderr)
        
        # Output JSON
        print(json.dumps(valid_questions, ensure_ascii=False))
        sys.exit(0)
        
    except Exception as e:
        print(json.dumps({"error": str(e)}))
        sys.exit(1)

if __name__ == "__main__":
    main()
