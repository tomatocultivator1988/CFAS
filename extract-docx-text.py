#!/usr/bin/env python3
"""
Simple DOCX text extractor
Extracts ALL text from a DOCX file and outputs to stdout
"""
import sys
from docx import Document

def extract_text(docx_path):
    """Extract all text from DOCX file"""
    try:
        doc = Document(docx_path)
        all_text = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                all_text.append(text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    all_text.append('\t'.join(row_text))
        
        return '\n'.join(all_text)
    except Exception as e:
        print(f"ERROR: {str(e)}", file=sys.stderr)
        sys.exit(1)

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python extract-docx-text.py <docx_file>", file=sys.stderr)
        sys.exit(1)
    
    docx_file = sys.argv[1]
    text = extract_text(docx_file)
    print(text)
