#!/usr/bin/env python3
"""
Inspect DOCX structure to understand Q1-Q2 format
"""
import sys
from docx import Document

docx_file = sys.argv[1] if len(sys.argv) > 1 else 'Aquaculture_set A.docx'

doc = Document(docx_file)

print("=" * 60)
print("  DOCX STRUCTURE INSPECTION")
print("=" * 60)
print()

# Check first few paragraphs
print("[PARAGRAPHS]")
for i, para in enumerate(doc.paragraphs[:15]):
    text = para.text.strip()
    if text:
        print(f"Para {i}: {text[:80]}")

print()
print("[TABLES]")
# Check first table
if doc.tables:
    table = doc.tables[0]
    print(f"First table: {len(table.rows)} rows x {len(table.columns)} columns")
    print()
    
    # Show first 5 rows
    for i, row in enumerate(table.rows[:5]):
        print(f"Row {i}:")
        for j, cell in enumerate(row.cells):
            text = cell.text.strip()
            if text:
                print(f"  Cell {j}: {text[:60]}")
        print()

print("=" * 60)
