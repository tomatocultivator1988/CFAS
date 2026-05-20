"""
Parse DOCX with Answer Key at Bottom
Handles format: Questions 1-100, then Answer Key table at bottom

No AI required - uses smart regex parsing
"""

import sys
import json
import re
from docx import Document

def extract_answer_key_from_table(doc):
    """Extract answer key from table at bottom of document"""
    answer_key = {}
    
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) >= 2:
                # Try to parse: "1" | "c" or "1." | "c" format
                cell1 = cells[0].text.strip()
                cell2 = cells[1].text.strip()
                
                # Extract question number
                q_match = re.match(r'(\d+)', cell1)
                if q_match:
                    q_num = int(q_match.group(1))
                    
                    # Extract answer letter
                    a_match = re.match(r'([a-d])', cell2, re.IGNORECASE)
                    if a_match:
                        answer_letter = a_match.group(1).lower()
                        answer_key[q_num] = answer_letter
    
    return answer_key

def extract_answer_key_from_text(text_lines):
    """Extract answer key from text format (not in table)"""
    answer_key = {}
    
    # Look for patterns like:
    # "1. c" or "1) c" or "1 - c" or "1: c"
    for line in text_lines:
        matches = re.findall(r'(\d+)[\.\)\-:]\s*([a-d])', line, re.IGNORECASE)
        for match in matches:
            q_num = int(match[0])
            answer_letter = match[1].lower()
            answer_key[q_num] = answer_letter
    
    return answer_key

def parse_questions(doc, answer_key):
    """Parse questions from document paragraphs"""
    questions = []
    
    current_question = None
    current_question_num = 0
    current_choices = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Check if this is a question (starts with number)
        q_match = re.match(r'^(\d+)\.\s*(.+)', text)
        if q_match:
            # Save previous question
            if current_question and current_choices:
                questions.append({
                    'question_text': current_question,
                    'answer_choices': current_choices
                })
            
            # Start new question
            current_question_num = int(q_match.group(1))
            current_question = q_match.group(2).strip()
            current_choices = []
            continue
        
        # Check if this is a choice (starts with letter)
        c_match = re.match(r'^([a-d])[\.\)]\s*(.+)', text, re.IGNORECASE)
        if c_match and current_question:
            choice_letter = c_match.group(1).lower()
            choice_text = c_match.group(2).strip()
            
            # Check if this is the correct answer from answer key
            correct_letter = answer_key.get(current_question_num, '').lower()
            is_correct = (choice_letter == correct_letter)
            
            current_choices.append({
                'choice_text': choice_text,
                'is_correct': is_correct
            })
            continue
        
        # If not a question or choice, might be continuation of question text
        if current_question and not current_choices:
            # Append to current question
            current_question += " " + text
    
    # Don't forget last question
    if current_question and current_choices:
        questions.append({
            'question_text': current_question,
            'answer_choices': current_choices
        })
    
    return questions

def validate_questions(questions):
    """Validate and fix questions"""
    valid_questions = []
    
    for q in questions:
        # Must have question text
        if not q.get('question_text'):
            continue
        
        # Must have at least 2 choices
        if not q.get('answer_choices') or len(q['answer_choices']) < 2:
            continue
        
        # Must have at least one correct answer
        has_correct = any(c.get('is_correct') for c in q['answer_choices'])
        if not has_correct:
            # Default first choice as correct
            q['answer_choices'][0]['is_correct'] = True
        
        # Ensure only one correct answer
        correct_count = sum(1 for c in q['answer_choices'] if c.get('is_correct'))
        if correct_count > 1:
            # Keep only first correct answer
            found_first = False
            for c in q['answer_choices']:
                if c.get('is_correct'):
                    if found_first:
                        c['is_correct'] = False
                    else:
                        found_first = True
        
        valid_questions.append(q)
    
    return valid_questions

def main():
    if len(sys.argv) < 2:
        print(json.dumps({"error": "Usage: python parse-docx-with-answerkey.py <docx-file>"}))
        sys.exit(1)
    
    docx_file = sys.argv[1]
    
    try:
        # Load document
        doc = Document(docx_file)
        
        # Step 1: Extract answer key from table
        answer_key = extract_answer_key_from_table(doc)
        
        # Step 2: If no table found, try to extract from text
        if not answer_key:
            text_lines = [para.text for para in doc.paragraphs]
            answer_key = extract_answer_key_from_text(text_lines)
        
        # Step 3: Parse questions
        questions = parse_questions(doc, answer_key)
        
        # Step 4: Validate questions
        valid_questions = validate_questions(questions)
        
        if not valid_questions:
            print(json.dumps({"error": "No valid questions found"}))
            sys.exit(1)
        
        # Output statistics to stderr (for debugging)
        print(f"Found {len(answer_key)} answers in answer key", file=sys.stderr)
        print(f"Parsed {len(valid_questions)} valid questions", file=sys.stderr)
        
        # Output JSON to stdout
        print(json.dumps(valid_questions, ensure_ascii=False))
        sys.exit(0)
        
    except Exception as e:
        print(json.dumps({"error": str(e)}))
        sys.exit(1)

if __name__ == "__main__":
    main()
