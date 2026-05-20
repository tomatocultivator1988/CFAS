#!/usr/bin/env python3
"""
Extract Cap_set...E.docx to see the answer key
"""
from docx import Document

try:
    # Read the file
    doc = Document('Exam-Main/Cap_set...E.docx')
    
    # Extract all text
    full_text = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text.append(text)
    
    # Save to file
    with open('Exam-Main/capset_e_raw.txt', 'w', encoding='utf-8') as f:
        f.write('\n'.join(full_text))
    
    # Also check for tables (answer key might be in a table)
    print("📋 Checking for tables...")
    table_data = []
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(' | '.join(row_data))
    
    if table_data:
        with open('Exam-Main/capset_e_tables.txt', 'w', encoding='utf-8') as f:
            f.write('\n'.join(table_data))
        print(f"✅ Found {len(doc.tables)} table(s)")
        print("📄 Table data saved to: Exam-Main/capset_e_tables.txt")
    
    print("✅ Text extracted to: Exam-Main/capset_e_raw.txt")
    print("")
    print("📝 Checking last 50 lines for answer key...")
    print("=" * 60)
    for line in full_text[-50:]:
        print(line)
    
except Exception as e:
    print(f"❌ Error: {e}")
