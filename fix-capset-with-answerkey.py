#!/usr/bin/env python3
"""
Fix Cap_set...F.docx and add the answer key properly
"""
from docx import Document

# Answer key data from your original message
answer_key = {
    1: 'd', 2: 'b', 3: 'b', 4: 'd', 5: 'b', 6: 'a', 7: 'd', 8: 'b', 9: 'd', 10: 'b',
    11: 'a', 12: 'd', 13: 'b', 14: 'c', 15: 'a', 16: 'b', 17: 'a', 18: 'b', 19: 'c', 20: 'a',
    21: 'a', 22: 'b', 23: 'c', 24: 'a', 25: 'b', 26: 'a', 27: 'b', 28: 'd', 29: 'c', 30: 'b',
    31: 'a', 32: 'd', 33: 'd', 34: 'b', 35: 'b', 36: 'a', 37: 'c', 38: 'a', 39: 'd', 40: 'b',
    41: 'c', 42: 'c', 43: 'b', 44: 'a', 45: 'c', 46: 'a', 47: 'b', 48: 'b', 49: 'd', 50: 'a',
    51: 'b', 52: 'a', 53: 'c', 54: 'd', 55: 'c', 56: 'd', 57: 'd', 58: 'a', 59: 'd', 60: 'b',
    61: 'a', 62: 'b', 63: 'a', 64: 'd', 65: 'd', 66: 'a', 67: 'c', 68: 'a', 69: 'a', 70: 'a',
    71: 'd', 72: 'c', 73: 'a', 74: 'a', 75: 'c', 76: 'b', 77: 'a', 78: 'a', 79: 'a', 80: 'd',
    81: 'b', 82: 'c', 83: 'b', 84: 'b', 85: 'd', 86: 'c', 87: 'a', 88: 'b', 89: 'c', 90: 'a',
    91: 'd', 92: 'd', 93: 'a', 94: 'd', 95: 'c', 96: 'd', 97: 'a', 98: 'a', 99: 'a', 100: 'b'
}

# Read the original file
doc = Document('Exam-Main/Cap_set...F.docx')

# Create new document
new_doc = Document()
new_doc.add_heading('CapSet F - Fisheries Examination', 0)

# Process paragraphs - split choices onto separate lines
import re

for para in doc.paragraphs:
    text = para.text.strip()
    
    if not text or 'Answer Key' in text or 'CapSet' in text:
        continue
    
    # Check if this line has multiple choices (tabs or multiple spaces)
    if re.search(r'[a-d]\.\s+.+[\t\s]{2,}[a-d]\.\s+', text):
        # Split by tabs or 2+ spaces before choice letters
        parts = re.split(r'[\t]+|\s{2,}(?=[a-d]\.\s+)', text)
        
        for part in parts:
            part = part.strip()
            if part:
                new_doc.add_paragraph(part)
    else:
        # Regular line
        new_doc.add_paragraph(text)

# Add answer key section
new_doc.add_page_break()
new_doc.add_heading('ANSWER KEY', 1)

# Add answer key as a table
table = new_doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'

# Header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Question'
hdr_cells[1].text = 'Answer'

# Add all answers
for q_num in range(1, 101):
    row_cells = table.add_row().cells
    row_cells[0].text = str(q_num)
    row_cells[1].text = answer_key[q_num].upper()

# Save
new_doc.save('Exam-Main/Cap_set_F_Complete.docx')

print("✅ Complete file created: Exam-Main/Cap_set_F_Complete.docx")
print("")
print("📋 Includes:")
print("   - All 100 questions (properly formatted)")
print("   - Each choice on its own line")
print("   - Answer key table at the end")
print("")
print("📤 Upload this file to import!")
