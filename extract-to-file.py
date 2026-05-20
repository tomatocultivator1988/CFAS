#!/usr/bin/env python3
"""
DOCX Text Extractor - Saves to file
Extracts ALL text from DOCX and saves to output file
"""
import sys
import os
from docx import Document

def extract_text_to_file(docx_path, output_path):
    """Extract all text from DOCX and save to file"""
    try:
        doc = Document(docx_path)
        all_text = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                all_text.append(text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    all_text.append('\t'.join(row_text))
        
        # Write to output file
        full_text = '\n'.join(all_text)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(full_text)
        
        # Print stats to stdout
        print(f"SUCCESS: Extracted {len(full_text)} characters, {len(all_text)} lines")
        return 0
        
    except Exception as e:
        print(f"ERROR: {str(e)}", file=sys.stderr)
        return 1

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python extract-to-file.py <docx_file> <output_file>", file=sys.stderr)
        sys.exit(1)
    
    docx_file = sys.argv[1]
    output_file = sys.argv[2]
    
    if not os.path.exists(docx_file):
        print(f"ERROR: File not found: {docx_file}", file=sys.stderr)
        sys.exit(1)
    
    exit_code = extract_text_to_file(docx_file, output_file)
    sys.exit(exit_code)
