#!/usr/bin/env python3
"""
Reformat Cap_set...F.docx to proper format for AI import
"""
import sys
from docx import Document

try:
    # Read the original file
    doc = Document('Exam-Main/Cap_set...F.docx')
    
    # Create new document
    new_doc = Document()
    new_doc.add_heading('CapSet F - Fisheries Examination', 0)
    
    # Extract all text
    full_text = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text.append(text)
    
    # Join and split by question numbers
    content = '\n'.join(full_text)
    
    # Save extracted text for inspection
    with open('Exam-Main/capset_f_raw.txt', 'w', encoding='utf-8') as f:
        f.write(content)
    
    print("✅ Extracted text saved to: Exam-Main/capset_f_raw.txt")
    print("📝 Please check the file to see the current format")
    print("")
    print("The AI import should work with the answer key table at the end!")
    print("Just upload the file directly to the admin panel.")
    
except Exception as e:
    print(f"❌ Error: {e}")
    print("")
    print("Make sure python-docx is installed:")
    print("pip install python-docx")
